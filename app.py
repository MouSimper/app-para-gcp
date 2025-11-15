# NOMBRE DEL ARCHIVO: app.py
# (Versión corregida para Cloud Run CPU + API de Gemini)

import os, sys, traceback, re, unicodedata
import torch
print(f"CUDA disponible: {torch.cuda.is_available()}") # Imprimirá False, es normal
import torch.nn.functional as F
import numpy as np
import gradio as gr
from datetime import datetime
from transformers import pipeline
from langdetect import detect, LangDetectException
from deep_translator import GoogleTranslator

from transformers import (
    BertTokenizer, BertForSequenceClassification,
    RobertaTokenizer, RobertaForSequenceClassification
)
from huggingface_hub import login as hf_login, HfFolder
import joblib

# --- API de Gemini (Reemplaza a Llama-3) ---
import vertexai
from vertexai.generative_models import GenerativeModel

try:
    # Inicializa Vertex AI (usa tu proyecto y región)
    # Cloud Run define automáticamente estas variables de entorno
    PROJECT_ID = os.environ.get("GOOGLE_CLOUD_PROJECT", "master-anagram-406222")
    LOCATION = "us-central1"
    vertexai.init(project=PROJECT_ID, location=LOCATION)
    
    # Carga el modelo Gemini
    gemini_model = GenerativeModel(model_name="gemini-1.5-flash-001")
    print("Modelo Gemini (Vertex AI) inicializado.")
except Exception as e:
    print(f"ERROR: No se pudo inicializar Vertex AI: {e}")
    gemini_model = None

# ---------- DOCX requerido ----------
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    print("Librería python-docx cargada.")
except Exception:
    print("\n[ERROR] Falta 'python-docx'. Instálalo con:\n   pip install python-docx\n")
    # No salimos, solo imprimimos

# =============== Control de Modo ===============
_mode_state = {"mode": "chat"}

def _detect_intent_for_mode(user_input: str) -> str:
    if not isinstance(user_input, str): return "chat"
    t = user_input.lower()
    if re.search(r"\b(crear|generar|hacer|elaborar).*(reporte|informe|documento)\b", t):
        return "create_report"
    if re.search(r"\b(cancelar|no quiero|dejar|salir|no generar|no hagas)\b", t):
        return "cancel_report"
    return "chat"

# =============== Token HF ===============
HF_TOKEN = os.environ.get("HF_TOKEN") or HfFolder.get_token()
if HF_TOKEN:
    try:
        hf_login(HF_TOKEN)
        HfFolder.save_token(HF_TOKEN)
        print("Hugging Face: autenticado desde variable de entorno o caché.")
    except Exception as e:
        print(f"Advertencia HF: {e}")
else:
    print("⚠️ No se detectó HF_TOKEN. Los modelos se descargarán anónimamente.")

# --- Traductor español → inglés ---
try:
    print("Cargando traductor Helsinki-NLP...")
    translator = pipeline("translation", model="Helsinki-NLP/opus-mt-es-en")
    print("Traductor cargado.")
except Exception as e:
    print(f"Error cargando traductor Helsinki: {e}")
    translator = None

def translate_to_english(text):
    if not isinstance(text, str) or not text.strip() or not translator:
        return text
    
    t = text.strip()
    try:
        lang = detect(t)
    except LangDetectException:
        lang = None
    
    spanish_markers = [" la ", " el ", "puerta", "cierra", "defecto", "ajuste"]
    should_translate = (
        lang == "es"
        or (lang is None and any(m in t.lower() for m in spanish_markers))
        or (not all(ord(c) < 128 for c in t))
    )
    if should_translate:
        try:
            translation = translator(t)[0]['translation_text']
            print(f"[Translator] '{t}' → '{translation}'")
            return translation
        except Exception as e:
            print(f"[Translator error] {e} ; returning original text")
            return t
    return t

# =============== Clasificador (CPU) ===============
DEVICE_CLF = torch.device("cpu")
BERT_CKPT     = "bert-finetuned-epoch7-acc0.9812"
ROBERTA_CKPT  = "roberta-finetuned-epoch6-acc0.9812"
LABEL_ENCODER = "label_encoder.pkl"
MAX_LEN_CLF   = 128
BEST_W        = 0.50
CLASS_NAMES = []
LABELS_ES_MAP = {
    "scratch": "rayadura", "dent": "abolladura", "corrosion": "corrosión",
    "crack": "grieta", "discoloration": "decoloración", "burr": "rebaba",
    "impact": "impacto", "abrasion": "abrasión", "handling": "manejo",
    "transport": "transporte", "manufacturing": "fabricación", "others":"otros",
}
bert = roberta = bert_tok = rob_tok = None

def _to_spanish_label(label_en: str) -> str:
    return LABELS_ES_MAP.get(str(label_en).lower(), str(label_en))

def _try_load_classifier():
    global bert, roberta, bert_tok, rob_tok, CLASS_NAMES
    try:
        # Plan A: Cargar localmente (para eso subimos el .pkl)
        print(f"Intentando cargar '{LABEL_ENCODER}' localmente...")
        label_encoder = joblib.load(LABEL_ENCODER)
        CLASS_NAMES = list(label_encoder.classes_)
        print(f"'{LABEL_ENCODER}' cargado. Clases: {len(CLASS_NAMES)}")
    except Exception as e_pkl:
        print(f"ERROR: No se pudo cargar '{LABEL_ENCODER}': {e_pkl}")
        return False # Si no hay .pkl, no podemos clasificar

    try:
        # Plan B: Descargar modelos de HF
        print(f"Intentando descargar clasificador BERT: {BERT_CKPT}...")
        bert_tok = BertTokenizer.from_pretrained(BERT_CKPT, token=HF_TOKEN)
        bert = BertForSequenceClassification.from_pretrained(BERT_CKPT, token=HF_TOKEN).to(DEVICE_CLF)
        print(f"Intentando descargar clasificador RoBERTa: {ROBERTA_CKPT}...")
        rob_tok = RobertaTokenizer.from_pretrained(ROBERTA_CKPT, token=HF_TOKEN)
        roberta = RobertaForSequenceClassification.from_pretrained(ROBERTA_CKPT, token=HF_TOKEN).to(DEVICE_CLF)
        print("Clasificador híbrido cargado (descargado, CPU).")
        return True
    except Exception as e_remote:
        print(f"ERROR: Error descargando modelos clasificadores: {e_remote}")
        return False

_classifier_ok = _try_load_classifier()

def _clean_for_clf(s: str) -> str:
    if not isinstance(s, str): s = str(s) if s is not None else ""
    s = unicodedata.normalize("NFKC", s)
    s = re.sub(r"[^\S\r\n]+", " ", s)
    return s.strip()

def _translate_if_needed(text: str) -> str:
    try:
        if not text.strip(): return text
        translated = GoogleTranslator(source='auto', target='en').translate(text)
        if translated and translated.lower() != text.lower():
            print(f"[DEBUG] Texto traducido automáticamente: {translated}")
        return translated
    except Exception as e:
        print(f"[WARN] No se pudo traducir: {e}")
        return text

@torch.no_grad()
def classify_text_hybrid(text: str, best_w: float = BEST_W):
    text = _clean_for_clf(text)
    text = _translate_if_needed(text)
    if not _classifier_ok:
        print("ERROR: classify_text_hybrid llamado pero _classifier_ok es False.")
        return {"error": "Clasificador no cargado."}
    try:
        eb = bert_tok(text, return_tensors="pt", truncation=True, padding=True, max_length=MAX_LEN_CLF).to(DEVICE_CLF)
        er =  rob_tok(text, return_tensors="pt", truncation=True, padding=True, max_length=MAX_LEN_CLF).to(DEVICE_CLF)
        pb = F.softmax(bert(**eb).logits, dim=1)
        pr = F.softmax(roberta(**er).logits, dim=1)
        probs = (best_w * pb + (1 - best_w) * pr)[0].detach().cpu().numpy()
        pred_idx = int(np.argmax(probs))
        label_en = CLASS_NAMES[pred_idx]
        label_es = _to_spanish_label(label_en)
        return {
            "label": label_en, "label_es": label_es,
            "confidence": float(probs[pred_idx] * 100),
        }
    except Exception as e:
        print(f"ERROR: Falla en inferencia de clasificación: {e}")
        return {"error": f"Falla en inferencia: {e}"}


# =============== Checklists SÍ/NO ===============
CHK_MARCO = [
    ("chk_m_unic_dim", "Unicación y dimensiones de vano terminado es conforme"),
    ("chk_m_nivel_piso", "Nivel de Piso terminado y sin desnivel"),
    ("chk_m_armado_marco", "Armado de piezas de marco de puerta (a nivel de piso) es conforme"),
    ("chk_m_sin_defectos", "No presenta defectos físicos (abolladuras, quiñes, etc.)"),
    ("chk_m_fijacion", "Perforaciones y fijación de marco con tirafones"),
    ("chk_m_verticalidad", "La instalación cumple con los requerimientos de verticalidad y acabado"),
    ("chk_m_otros", "Otros _____"),
]
CHK_HOJA = [
    ("chk_h_aplomado", "Marco de puerta está aplomado a nivel"),
    ("chk_h_medidas_planos", "Medidas nominales de la hoja de puerta están de acuerdo a los Planos del Proyecto"),
    ("chk_h_inst_hoja", "Correcta instalación de hoja de puerta (incluyen bisagras)"),
    ("chk_h_cierre_juego", "Cierre y juego adecuado"),
    ("chk_h_acabado", "Acabado superficial conforme"),
    ("chk_h_limpieza", "Limpieza final del área de trabajo"),
    ("chk_h_otros", "Otros _____"),
]

# =============== Plantillas ===============
TEMPLATES = {
    "informe_incidente": {
        "title": "Informe de Incidente",
        "fields": [
            {"key":"cliente","label":"Cliente", "required":False},
            {"key":"contratista","label":"Contratista", "required":False},
            {"key":"subcontratista","label":"Subcontratista", "required":False},
            {"key":"piso_sector","label":"Piso / Sector", "required":False},
            {"key":"supervision","label":"Supervisión", "required":False},
            {"key":"nro_protocolo","label":"Número de protocolo", "required":False},
            {"key": "fecha", "label": "Fecha del incidente", "required": True, "hint":"Ej: 08/10/2025"},
            {"key": "lugar", "label": "Lugar", "required": True, "hint":"Ej: Planta A - Línea 3"},
            {"key": "ubicacion","label":"Ubicación de la puerta", "required":False},
            {"key": "responsable", "label": "Responsable", "required": True, "hint":"Nombre y cargo"},
            {"key": "acciones", "label": "Acciones correctivas", "required": False, "hint":"Medidas aplicadas y plan"},
            *[{"key": k, "label": f"(Marco) {lbl}", "required": False, "hint": "Responde: si / no"} for k, lbl in CHK_MARCO],
            *[{"key": k, "label": f"(Hoja) {lbl}", "required": False, "hint": "Responde: si / no"} for k, lbl in CHK_HOJA],
            {"key": "observaciones", "label": "Observaciones", "required": False, "hint":"Notas a clasificar"},
        ],
    },
}

# =============== Helpers de intents naturales ===============
def _normalize(s: str) -> str:
    s = s.lower().strip()
    s = ''.join(c for c in unicodedata.normalize('NFD', s) if unicodedata.category(c) != 'Mn')
    return s
YES_SET = {"si","sí","s","x","true","1","yes"}
NO_SET  = {"no","n","false","0"}

def _parse_yesno(value: str):
    if value is None: return None
    t = _normalize(str(value))
    if t in YES_SET: return "si"
    if t in NO_SET: return "no"
    if re.search(r"\bsi\b|\bsí\b", t): return "si"
    if re.search(r"\bno\b", t): return "no"
    return None

def _match_template(user_text_norm: str):
    alias = {"informe_incidente": ["informe de incidente", "reporte de incidente", "informe incidente", "incidente", "protocolo de instalación"]}
    for key, patterns in alias.items():
        if any(p in user_text_norm for p in patterns):
            return key
    for k, meta in TEMPLATES.items():
        if _normalize(meta["title"]) in user_text_norm or k in user_text_norm:
            return k
    return None

def route_intent(user_text: str, state: dict):
    t = _normalize(user_text)
    if any(w in t for w in ["plantilla", "plantillas", "ayuda", "opciones", "menu"]):
        return {"type":"list_templates", "payload":{}}
    if any(w in t for w in ["cancelar", "anular", "reiniciar", "reset"]):
        return {"type":"cancel", "payload":{}}
    if any(w in t for w in ["descarga", "descargar", "generar archivo", "documento"]):
        return {"type":"download", "payload":{}}
    tk = _match_template(t)
    if tk:
        return {"type":"start_template", "payload":{"template_key": tk}}
    if state and state.get("active") and not state.get("finished"):
        return {"type":"answer_field", "payload":{"text": user_text}}
    return {"type":"chitchat", "payload":{"text": user_text}}

# =============== Generación LLM (API de Gemini) ===============
def llama_chat_generate(user_message: str, history_pairs: list | None = None) -> str:
    history_pairs = history_pairs or []
    
    # Prepara el historial para Gemini
    gemini_history = []
    gemini_history.append({"role": "user", "parts": [{"text": (
        "Eres un asistente que guía al usuario para crear documentos a partir de plantillas y "
        "clasificar automáticamente la sección 'Observaciones'. Si el usuario parece querer iniciar "
        "un 'Informe de Incidente', sugiere iniciar la plantilla correspondiente, "
        "preguntando el primer campo de forma directa y breve."
    )}]})
    gemini_history.append({"role": "model", "parts": [{"text": "Entendido. Estoy listo para ayudar a crear el 'Informe de Incidente'."}]})

    for u, a in history_pairs:
        gemini_history.append({"role": "user", "parts": [{"text": u}]})
        gemini_history.append({"role": "model", "parts": [{"text": a}]})
        
    try:
        if not gemini_model:
            print("ERROR: gemini_model no está inicializado.")
            raise Exception("El modelo Gemini no se inicializó correctamente.")
        
        # Crea una sesión de chat y envía el mensaje
        chat_session = gemini_model.start_chat(history=gemini_history)
        response = chat_session.send_message(user_message)
        
        # Devuelve solo el texto de la respuesta
        return response.text
        
    except Exception as e:
        print(f"[ERROR en API de Gemini] {e}")
        traceback.print_exc()
        return f"Hubo un error al contactar la API de Gemini: {e}"

# =============== Wizard por chat (estado) ===============
def _reset_session():
    return {
        "active": False, "template_key": None, "fields": [], "idx": 0, "answers": {},
        "cls_result": None, "file_path": None, "finished": False, "custom_collect": None,
    }

def _render_next_question(state):
    f = state["fields"][state["idx"]]
    if f["key"] in {"chk_m_otros", "chk_h_otros"}:
        prompt = _start_custom_other_block(state, f) if not state.get("custom_collect") else (
            f"Continuemos con **{f['label']}**. Escribe la descripción o responde **no** para avanzar."
        )
        return prompt
    req = " (obligatorio)" if f.get("required") else ""
    hint = f.get("hint") or ""
    return f"**{f['label']}{req}:**\n_{hint}_"

CUSTOM_OTHER_KEYS = {"chk_m_otros", "chk_h_otros"}

def _start_custom_other_block(state, field):
    prompt = (
        f"En el apartado **{field['label']}** puedes agregar descripciones adicionales. "
        "Escribe el texto del punto extra o responde **no** para omitir."
    )
    state["custom_collect"] = {
        "field_key": field["key"], "label": field["label"], "entries": [],
        "step": "ask_desc", "current_desc": None,
    }
    return prompt

def _handle_custom_collect(state, user_msg):
    data = state.get("custom_collect")
    if not data: return None
    text = (user_msg or "").strip()
    if data["step"] == "ask_desc":
        if not text:
            return {"message": "Escribe una descripción o responde **no** para continuar.", "advance": False}
        if _normalize(text) == "no":
            state["answers"][data["field_key"]] = list(data["entries"])
            state["custom_collect"] = None
            message = ("No se agregarán más puntos en 'Otros'." if data["entries"] else "Sin puntos adicionales en 'Otros'.")
            return {"message": message, "advance": True}
        data["current_desc"] = text
        data["step"] = "ask_mark"
        return {"message": f"¿Marcamos \"{text}\" como **sí** o **no**?", "advance": False}
    if data["step"] == "ask_mark":
        choice = _parse_yesno(text)
        if choice is None:
            return {"message": "Responde solo **sí** o **no** para ese punto.", "advance": False}
        data["entries"].append({"text": data["current_desc"], "value": choice})
        data["current_desc"] = None
        data["step"] = "ask_desc"
        return {"message": "¿Deseas agregar otro punto en 'Otros'? Escribe la descripción o responde **no**.", "advance": False}
    return {"message": "", "advance": False}

# ============== Helpers de diseño DOCX ==============
LOGO_PATH = "./logo.png"
ACCENT_HEX = "1F3B73"
PALE_ACCENT = "ECF0FF"
LIGHT_BG = "F8FAFF"
ACCENT_RGB = RGBColor(31, 59, 115)
WHITE_RGB = RGBColor(255, 255, 255)

def _apply_base_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"; normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.2
    normal.paragraph_format.space_after = Pt(4)
    
def _set_cell_shading(cell, fill_hex="D9D9D9"):
    try:
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill_hex); tcPr.append(shd)
    except Exception as e:
        print(f"Error aplicando shading: {e}")

def _cell_p(cell, text="", bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=10, color=None):
    p = cell.paragraphs[0]; p.alignment = align
    run = p.add_run(text); run.bold = bold; run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    return p

def _add_spacer(doc: Document, h=0.12):
    doc.add_paragraph().paragraph_format.space_after = Pt(h*72)

def _section_heading(doc: Document, title: str, subtitle: str = ""):
    tbl = doc.add_table(rows=1, cols=1); cell = tbl.rows[0].cells[0]
    _set_cell_shading(cell, PALE_ACCENT)
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title.upper()); run.bold = True
    run.font.size = Pt(11.5); run.font.color.rgb = ACCENT_RGB
    _add_spacer(doc, 0.05)
    return tbl

def _card_block(doc: Document, title: str, body: str):
    tbl = doc.add_table(rows=1, cols=1); cell = tbl.rows[0].cells[0]
    _set_cell_shading(cell, LIGHT_BG)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell.paragraphs[0].add_run(title).bold = True
    cell.add_paragraph(body or "—")
    _add_spacer(doc, 0.08)
    return tbl

def _signature_block(doc: Document):
    entries = [
        ("Inspector responsable", "________________________    Fecha: ____________"),
        ("Responsable del proyecto", "________________________    Fecha: ____________"),
    ]
    tbl = doc.add_table(rows=len(entries), cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, line) in enumerate(entries):
        tbl.cell(i, 0).text = label; tbl.cell(i, 1).text = line
    _add_spacer(doc, 0.1)
    return tbl

def _kv_table(doc, headers_and_values, cols=4):
    rows = (len(headers_and_values) + cols//2) // (cols//2)
    tbl = doc.add_table(rows=rows, cols=cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    r_i = c_i = 0
    for k, v in headers_and_values:
        label_cell = tbl.cell(r_i, c_i); value_cell = tbl.cell(r_i, c_i+1)
        _set_cell_shading(label_cell, "E3E8F7"); _set_cell_shading(value_cell, "FFFFFF")
        _cell_p(label_cell, k, bold=True, size=10, color=ACCENT_RGB)
        _cell_p(value_cell, v or "", size=10)
        c_i += 2
        if c_i >= cols: c_i = 0; r_i += 1
    _add_spacer(doc, 0.10)
    return tbl

def _mark_for_yes(v: str) -> str:
    return "X" if _parse_yesno(v) == "si" else ""

def _mark_for_no(v: str) -> str:
    return "X" if _parse_yesno(v) == "no" else ""

def _clean_check_label(label: str) -> str:
    return re.sub(r"^\((Marco|Hoja)\)\s*", "", label).strip()

def _expand_checklist_items(items_with_keys, answers):
    rows = []
    for key, label in items_with_keys:
        if key in CUSTOM_OTHER_KEYS:
            extras = answers.get(key, [])
            if not isinstance(extras, list) or not extras: continue
            for idx, entry in enumerate(extras, start=1):
                desc = (entry.get("text") or f"{label} #{idx}").strip()
                rows.append((desc, entry.get("value", "")))
        else:
            rows.append((_clean_check_label(label), answers.get(key, "")))
    return rows

def _items_check_table(doc, titulo_seccion, items_with_values):
    _section_heading(doc, titulo_seccion, "Marca con X la opción correspondiente")
    t_hdr = doc.add_table(rows=1, cols=3); t_hdr.alignment = WD_TABLE_ALIGNMENT.CENTER
    c0, c1, c2 = t_hdr.rows[0].cells
    _set_cell_shading(c0, ACCENT_HEX); _cell_p(c0, "DESCRIPCIÓN", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=11, color=WHITE_RGB)
    _set_cell_shading(c1, ACCENT_HEX); _cell_p(c1, "SÍ", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=11, color=WHITE_RGB)
    _set_cell_shading(c2, ACCENT_HEX); _cell_p(c2, "NO", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=11, color=WHITE_RGB)
    t = doc.add_table(rows=len(items_with_values), cols=3); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, val) in enumerate(items_with_values):
        if i % 2 == 0:
            _set_cell_shading(t.cell(i,0), "F4F6FB"); _set_cell_shading(t.cell(i,1), "F4F6FB"); _set_cell_shading(t.cell(i,2), "F4F6FB")
        _cell_p(t.cell(i,0), label, size=10)
        pc_si = t.cell(i,1).paragraphs[0]; pc_si.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_si = pc_si.add_run(_mark_for_yes(val)); run_si.font.size = Pt(12)
        pc_no = t.cell(i,2).paragraphs[0]; pc_no.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_no = pc_no.add_run(_mark_for_no(val)); run_no.font.size = Pt(12)
    _add_spacer(doc, 0.15)
    return t

# =============== Generación DOCX ===============
def _generate_doc(template_key, answers, cls_res):
    title_img = "PROTOCOLO DE INSTALACIÓN DE PUERTAS CORTAFUEGO"
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    # Cambia el directorio de salida para que funcione en Cloud Run (que es solo lectura)
    out_dir = "/tmp/outputs" if os.environ.get("PORT") else "outputs"
    os.makedirs(out_dir, exist_ok=True)
    fpath = os.path.join(out_dir, f"Protocolo_{ts}.docx")
    print(f"Generando documento en: {fpath}")

    try:
        doc = Document(); _apply_base_styles(doc)
        sec = doc.sections[0]
        sec.top_margin = Inches(0.5); sec.bottom_margin = Inches(0.5)
        sec.left_margin = Inches(0.6); sec.right_margin = Inches(0.5)

        hdr_tbl = doc.add_table(rows=1, cols=2); hdr_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        left, right = hdr_tbl.rows[0].cells
        if os.path.exists(LOGO_PATH):
            try: left.paragraphs[0].add_run().add_picture(LOGO_PATH, width=Inches(1.2))
            except Exception as e: print(f"Error cargando logo: {e}")
        pr = right.paragraphs[0]; pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = pr.add_run(title_img); r.bold = True; r.font.size = Pt(13)
        _add_spacer(doc, 0.05)

        if template_key == "informe_incidente":
            _section_heading(doc, "Datos generales del protocolo", "Identificación del cliente y obra")
            _kv_table(doc, [
                ("CLIENTE:", str(answers.get("cliente",""))),
                ("FECHA:", str(answers.get("fecha",""))),
                ("CONTRATISTA:", str(answers.get("contratista",""))),
                ("UBICACIÓN DE LA PUERTA:", str(answers.get("ubicacion",""))),
                ("SUBCONTRATISTA:", str(answers.get("subcontratista",""))),
                ("PISO/SECTOR:", str(answers.get("piso_sector",""))),
                ("SUPERVISIÓN:", str(answers.get("supervision",""))),
                ("NÚMERO DE PROTOCOLO:", str(answers.get("nro_protocolo",""))),
            ], cols=4)
            marco_vals = _expand_checklist_items(CHK_MARCO, answers)
            _items_check_table(doc, "Puntos de control - Marco", marco_vals)
            hoja_vals = _expand_checklist_items(CHK_HOJA, answers)
            _items_check_table(doc, "Puntos de control - Hoja", hoja_vals)
            obs_text = str(answers.get("observaciones","")).strip()
            if obs_text: _card_block(doc, "Observaciones reportadas", obs_text)
            acciones_text = str(answers.get("acciones","")).strip() or "No se registraron acciones correctivas."
            _card_block(doc, "Comentarios y acciones correctivas", acciones_text)
            if obs_text:
                if cls_res and "error" not in cls_res:
                    label = cls_res.get("label_es", cls_res.get("label",""))
                    cls_body = f"Etiqueta predicha: {label}\nConfianza: {cls_res['confidence']:.2f}%"
                    _card_block(doc, "Clasificación automática de observaciones", cls_body)
                else:
                    _card_block(doc, "Clasificación automática de observaciones", "Clasificador no disponible.")
            _section_heading(doc, "Firmas de conformidad", "Validación de responsables")
            _signature_block(doc)

        doc.save(fpath)
        print(f"Documento guardado: {fpath}")
        return fpath
    except Exception as e:
        print(f"ERROR: No se pudo generar el DOCX: {e}")
        traceback.print_exc()
        return None

# =============== Chat Handler ===============
def chat_handler(message, chat_history, session_state, file_comp, quick_action=None):
    chat_history = chat_history or []
    state = session_state or _reset_session()
    user_msg = (message or "").strip()
    if quick_action: user_msg = quick_action

    pairs, last_user = [], None
    for m in chat_history:
        if m.get("role") == "user": last_user = m.get("content")
        elif m.get("role") == "assistant" and last_user is not None:
            pairs.append((last_user, m.get("content"))); last_user = None

    def send(bot_text):
        if user_msg: chat_history.append({"role":"user","content":user_msg})
        chat_history.append({"role":"assistant","content":bot_text})
        return chat_history, "", state, gr.update()

    intent = route_intent(user_msg, state)
    intent_mode = _detect_intent_for_mode(user_msg)
    if intent_mode == "create_report": _mode_state["mode"] = "report"
    elif intent_mode == "cancel_report":
        _mode_state["mode"] = "chat"
        if state and state.get("active"): state = _reset_session()
        chat_history.append({"role":"user","content":user_msg})
        chat_history.append({"role":"assistant","content":"Entendido, salimos del modo de reporte."})
        return chat_history, "", state, gr.update()

    if _mode_state.get("mode") == "chat":
        if intent["type"] not in {"start_template", "list_templates", "download", "cancel"} and not (state and state.get("active")):
            try: reply = llama_chat_generate(user_msg, pairs)
            except Exception as e: reply = f"Error generando respuesta: {e}"
            if user_msg: chat_history.append({"role":"user","content":user_msg})
            chat_history.append({"role":"assistant","content":reply})
            return chat_history, "", state, gr.update()

    if intent["type"] == "list_templates":
        names = "\n".join([f"- **{meta['title']}**" for _,meta in TEMPLATES.items()])
        return send(f"Plantillas disponibles:\n{names}"), "", state, gr.update()
    if intent["type"] == "cancel":
        state = _reset_session()
        return send("He cancelado la sesión actual."), "", state, gr.update()
    if intent["type"] == "download":
        if not state["active"] or not state["finished"]:
            return send("Aún no has completado la plantilla."), "", state, gr.update()
        fpath = _generate_doc(state["template_key"], state["answers"], state["cls_result"])
        if fpath is None:
             return send("Error: No se pudo generar el documento DOCX."), "", state, gr.update()
        state["file_path"] = fpath
        return send(f"Documento generado: **{os.path.basename(fpath)}**"), "", state, gr.update(value=fpath)
    if intent["type"] == "start_template":
        key = intent["payload"]["template_key"]
        state = _reset_session(); state["active"] = True; state["template_key"] = key
        state["fields"] = TEMPLATES[key]["fields"]; state["idx"] = 0
        q = _render_next_question(state); _mode_state["mode"] = "report"
        return send(f"Perfecto, iniciaré **{TEMPLATES[key]['title']}**.\n{q}"), "", state, gr.update(value=None)

    if (intent["type"] == "answer_field") and state["active"] and not state["finished"]:
        f = state["fields"][state["idx"]]
        if state.get("custom_collect"):
            custom = _handle_custom_collect(state, user_msg)
            if not custom or not custom["advance"]:
                reply = custom["message"] if custom else "Necesito una respuesta válida."
                return send(reply), "", state, gr.update()
            state["idx"] += 1
            if state["idx"] < len(state["fields"]):
                q = _render_next_question(state)
                combo = f"{custom.get('message') or 'Continuemos.'}\n\n{q}"
                return send(combo), "", state, gr.update()
            state["finished"] = True
            return send(custom.get("message") or "¡Plantilla completa! Di **descargar**."), "", state, gr.update()
        if f.get("required") and not user_msg:
            return send(f"El campo **{f['label']}** es obligatorio."), "", state, gr.update()
        if f["key"].startswith("chk_"):
            yn = _parse_yesno(user_msg)
            if yn is None:
                return send(f"⚠️ Opción no válida. Usa solo **sí** o **no**."), "", state, gr.update()
            state["answers"][f["key"]] = yn
        else:
            state["answers"][f["key"]] = user_msg
        just_obs = (f["key"] == "observaciones")
        state["idx"] += 1
        if just_obs and user_msg.strip():
            res = classify_text_hybrid(user_msg)
            state["cls_result"] = res
            if "error" in res:
                print(f"Error de clasificador: {res['error']}")
        if state["idx"] < len(state["fields"]):
            q = _render_next_question(state)
            return send(q), "", state, gr.update()
        state["finished"] = True
        return send("¡Plantilla completa! Di **descargar** para generar el documento."), "", state, gr.update()

    try: reply = llama_chat_generate(user_msg, pairs)
    except Exception as e: reply = f"Error generando respuesta: {e}"
    if user_msg: chat_history.append({"role":"user","content":user_msg})
    chat_history.append({"role":"assistant","content":reply})
    return chat_history, "", state, gr.update()

# =============== UI Gradio ===============
# (Tu CSS DARK_CSS se omite aquí por brevedad, pero puedes pegarlo)
DARK_CSS = """
<style>
  :root { --radius-lg: 16px; --blur-bg: rgba(12, 18, 32, 0.78); }
  body {
    background: #050910;
    color: #e6e9ef;
    font-family: 'Inter', ui-sans-serif, system-ui, -apple-system, Segoe UI, Roboto, "Helvetica Neue", Arial, "Noto Sans";
    position: relative;
    overflow-x: hidden;
  }
  body::before {
    content: "";
    position: fixed;
    inset: 0;
    background: radial-gradient(circle at 20% 20%, rgba(79,111,255,0.15), transparent 40%),
                radial-gradient(circle at 80% 0%, rgba(147,70,255,0.15), transparent 45%),
                radial-gradient(circle at 50% 80%, rgba(23,188,255,0.12), transparent 40%);
    filter: blur(40px);
    animation: aurora 24s ease-in-out infinite alternate;
    pointer-events: none;
    z-index: -2;
  }
  .gradio-container {
    max-width: 1220px !important;
    margin: 0 auto !important;
    padding: 32px 28px 64px !important;
    background: radial-gradient(circle at top, rgba(49,91,255,0.18), transparent 55%) #050910;
  }
  .block, .group, .gr-panel, .gr-box, .form, .tabs, .tabitem, .tabs > div, .gradio-row {
    background: transparent !important;
    border: none !important;
    box-shadow: none !important;
  }
  .hero-card {
    background: var(--blur-bg);
    border: 1px solid rgba(255,255,255,0.08);
    border-radius: 24px;
    padding: 22px 26px;
    margin-bottom: 14px;
    backdrop-filter: blur(16px);
    box-shadow: 0 18px 38px rgba(0,0,0,0.45);
    opacity: 0;
    animation: floatUp 0.8s ease forwards;
  }
  .hero-card h1 {
    font-size: 28px;
    margin: 0;
    color: #f5f7ff;
  }
  .hero-sub {
    color: rgba(230,233,239,0.85);
    margin-top: 4px;
  }
  .hero-pill {
    display:inline-flex;
    align-items:center;
    gap:8px;
    padding:8px 14px;
    border-radius:999px;
    border:1px solid rgba(255,255,255,0.12);
    background: rgba(13,30,60,0.55);
    font-size:13px;
    animation: pulseGlow 3s ease-in-out infinite;
  }
  .main-layout {
    gap: 22px;
    align-items: stretch;
  }
  .panel {
    background: var(--blur-bg);
    border: 1px solid rgba(255,255,255,0.06);
    border-radius: 24px;
    padding: 18px 20px 22px;
    box-shadow: 0 10px 35px rgba(0,0,0,0.4);
    opacity: 0;
    animation: fadeIn 0.85s ease forwards;
  }
  .chat-panel .gradio-chatbot {
    background: rgba(7,12,22,0.9) !important;
    border: 1px solid rgba(255,255,255,0.04) !important;
    border-radius: 20px !important;
    box-shadow: inset 0 0 0 1px rgba(255,255,255,0.02), 0 15px 45px rgba(2,6,16,0.8);
    animation: floatDelay 1s ease forwards;
  }
  .chat-panel .gradio-chatbot > div {
    background: transparent !important;
  }
  .message.user, .chatbot .message.user {
    background: #111a29 !important;
    border-radius: 14px !important;
    border: 1px solid rgba(255,255,255,0.04) !important;
    animation: chatPop 0.35s ease;
  }
  .message.bot, .chatbot .message.bot {
    background: #0c131f !important;
    border-radius: 14px !important;
    border: 1px solid rgba(21,38,63,0.8) !important;
    animation: chatPop 0.35s ease;
  }
  .chat-panel .gr-textbox textarea {
    min-height: 54px !important;
  }
  .chat-panel .gr-textbox, textarea, .gr-input, input[type="text"] {
    background: rgba(8,13,23,0.9) !important;
    color: #f1f4ff !important;
    border: 1px solid rgba(255,255,255,0.08) !important;
    border-radius: 999px !important;
    padding: 14px 18px !important;
    font-size: 16px !important;
  }
  .chat-panel .chips-row {
    flex-wrap: wrap;
    gap: 10px;
    margin-bottom: 4px;
  }
  .chip-button {
    flex: 1 1 48%;
    min-width: 180px;
    background: rgba(25,47,87,0.7) !important;
    border: 1px solid rgba(65,110,255,0.35) !important;
    border-radius: 14px !important;
    color: #f0f3ff !important;
    font-weight: 600 !important;
    height: 48px !important;
    transition: transform .2s ease, box-shadow .2s ease, border .2s ease, filter .2s ease;
    animation: cascade 0.9s ease forwards;
  }
  .chip-button:nth-child(2n) { animation-delay: 0.15s; }
  .chip-button:nth-child(3n) { animation-delay: 0.3s; }
  .chip-button:nth-child(4n) { animation-delay: 0.45s; }
  .chip-button:hover {
    transform: translateY(-2px) scale(1.01);
    box-shadow: 0 12px 24px rgba(31,63,122,0.35);
    filter: brightness(1.05);
  }
  .file-card {
    margin-top: 12px;
    border-radius: 20px !important;
    background: rgba(10,17,30,0.9) !important;
    border: 1px dashed rgba(90,119,255,0.45) !important;
  }
  .info-panel .info-card {
    background: rgba(7,12,22,0.75);
    border: 1px solid rgba(255,255,255,0.05);
    border-radius: 18px;
    padding: 18px 20px;
    margin-bottom: 14px;
    opacity: 0;
    animation: slideIn 0.85s ease forwards;
  }
  .info-panel h3 {
    margin-top: 0;
    font-size: 18px;
    color: #f5f7ff;
  }
  .info-steps ol {
    padding-left: 22px;
    margin: 0;
    color: rgba(230,233,239,0.85);
  }
  .info-tags {
    display: flex;
    flex-wrap: wrap;
    gap: 8px;
  }
  .info-tags span {
    padding: 6px 12px;
    border-radius: 999px;
    border: 1px solid rgba(255,255,255,0.08);
    background: rgba(12,20,34,0.8);
    font-size: 13px;
  }
  .status-board {
    display: grid;
    grid-template-columns: repeat(auto-fit, minmax(140px, 1fr));
    gap: 10px;
  }
  .status-item {
    background: rgba(10,16,28,0.85);
    border: 1px solid rgba(255,255,255,0.05);
    border-radius: 16px;
    padding: 12px 14px;
    font-size: 13px;
    animation: cardPulse 2.4s ease-in-out infinite;
  }
  .status-item span {
    display: block;
    color: rgba(230,233,239,0.68);
    font-size: 12px;
  }
  @media (max-width: 980px) {
    .main-layout { flex-direction: column; }
    .chip-button { flex: 1 1 100%; }
  }
  @keyframes aurora {
    0% { transform: translate3d(-4%, -2%, 0) scale(1); }
    100% { transform: translate3d(6%, 4%, 0) scale(1.05); }
  }
  @keyframes floatUp {
    0% { opacity: 0; transform: translateY(20px); }
    100% { opacity: 1; transform: translateY(0); }
  }
  @keyframes fadeIn {
    0% { opacity: 0; transform: translateY(24px); }
    100% { opacity: 1; transform: translateY(0); }
  }
  @keyframes floatDelay {
    0% { opacity: 0; transform: translateY(18px); }
    100% { opacity: 1; transform: translateY(0); }
  }
  @keyframes slideIn {
    0% { opacity: 0; transform: translateX(30px); }
    100% { opacity: 1; transform: translateX(0); }
  }
  @keyframes cascade {
    0% { opacity: 0; transform: translateY(20px) scale(0.98); }
    100% { opacity: 1; transform: translateY(0) scale(1); }
  }
  @keyframes pulseGlow {
    0%, 100% { box-shadow: 0 0 10px rgba(98,130,255,0.25); }
    50% { box-shadow: 0 0 18px rgba(98,130,255,0.45); }
  }
  @keyframes chatPop {
    0% { opacity: 0; transform: translateY(8px) scale(0.98); }
    100% { opacity: 1; transform: translateY(0) scale(1); }
  }
  @keyframes cardPulse {
    0%, 100% { border-color: rgba(255,255,255,0.05); }
    50% { border-color: rgba(126,162,255,0.18); }
  }
</style>
"""

def build_and_launch_gradio():
    with gr.Blocks(title="Chat Plantillas (NFPA 80)", css=DARK_CSS, theme=gr.themes.Soft(
        primary_hue="indigo", neutral_hue="slate"
    )) as demo:

        gr.HTML("""
        <div class="hero-card">
          <div style="display:flex;justify-content:space-between;align-items:flex-start;gap:18px;flex-wrap:wrap;">
            <div>
              <h1>Asistente Inteligente NFPA 80</h1>
              <p class="hero-sub">Completa plantillas guiadas, obtiene una clasificación automática<br/>y genera documentos Word listos para entregar.</p>
            </div>
            <div class="hero-pill" id="status-pill">
              <span>Modo actual</span>
              <strong>Listo para empezar</strong>
            </div>
          </div>
        </div>
        """)

        with gr.Row(elem_classes=["main-layout"]):
            with gr.Column(scale=7, elem_classes=["panel", "chat-panel"]):
                chatbot = gr.Chatbot(height=520, type="messages", label=None)
                txt = gr.Textbox(
                    show_label=False,
                    placeholder="Haz una pregunta o di “Quiero un Informe de Incidente”...",
                    lines=1,
                    max_lines=1,
                    autofocus=True,
                )
                with gr.Row(elem_classes=["chips-row"]):
                    qa_create_inf = gr.Button("📝 Informe de incidente", elem_classes=["chip-button"])
                with gr.Row(elem_classes=["chips-row"]):
                    qa_list = gr.Button("📚 Ver plantillas", elem_classes=["chip-button"])
                    qa_download = gr.Button("⬇️ Descargar documento", elem_classes=["chip-button"])
                file_out = gr.File(label="Documento generado", interactive=False, elem_classes=["file-card"])
            with gr.Column(scale=5, elem_classes=["panel", "info-panel"]):
                gr.HTML("""
                <div class="info-card info-steps">
                  <h3>Cómo funciona</h3>
                  <ol>
                    <li>Inicia una plantilla con los botones rápidos o escribiendo “Quiero…”.</li>
                    <li>Responde las preguntas guiadas y los checklists sí/no.</li>
                    <li>El clasificador detecta la categoría de la observación.</li>
                    <li>Descarga el DOCX formateado con protocolo NFPA 80.</li>
                  </ol>
                </div>
                """)
                gr.HTML("""
                <div class="info-card">
                  <h3>Sugerencias rápidas</h3>
                  <div class="info-tags">
                    <span>Modo chat libre para dudas</span>
                    <span>Automatiza checklists</span>
                    <span>Traducción al inglés integrada</span>
                    <span>Clasificación híbrida BERT + RoBERTa</span>
                  </div>
                </div>
                """)
                state = gr.State(_reset_session())

        txt.submit(chat_handler, [txt, chatbot, state, file_out], [chatbot, txt, state, file_out])
        qa_create_inf.click(
            fn=chat_handler,
            inputs=[gr.Textbox(value="Quiero informe de incidente", visible=False), chatbot, state, file_out, gr.Textbox(value="Quiero informe de incidente", visible=False)],
            outputs=[chatbot, txt, state, file_out]
        )
        qa_list.click(
            fn=chat_handler,
            inputs=[gr.Textbox(value="Ver plantillas", visible=False), chatbot, state, file_out, gr.Textbox(value="Ver plantillas", visible=False)],
            outputs=[chatbot, txt, state, file_out]
        )
        qa_download.click(
            fn=chat_handler,
            inputs=[gr.Textbox(value="Descargar documento", visible=False), chatbot, state, file_out, gr.Textbox(value="Descargar documento", visible=False)],
            outputs=[chatbot, txt, state, file_out]
        )

    # ===> ¡CAMBIO IMPORTANTE PARA LA NUBE! <===
    print("Iniciando servidor Gradio en 0.0.0.0...")
    port = int(os.environ.get("PORT", 8080))
    demo.launch(server_name="0.0.0.0", server_port=port, share=False)

# =============== MAIN ===============
if __name__ == "__main__":
    try:
        build_and_launch_gradio()
    except Exception as e:
        print("Error lanzando Gradio:", e)
        traceback.print_exc(); sys.exit(1)